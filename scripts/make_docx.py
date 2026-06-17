#!/usr/bin/env python
"""Turn a finished markdown paper into an APA-7 formatted .docx (student paper).

Usage:
    python scripts/make_docx.py <input.md> [output.docx]

The input markdown may start with a simple front-matter block so the title page
is filled in correctly:

    ---
    title: Leadership Philosophy
    author: Jordan Reyes
    affiliation: Department of Human and Organizational Learning, The George Washington University
    course: HOL 6704: Leadership in Organizations
    instructor: Dr. M. Lewis
    date: June 17, 2026
    ---

    # Leadership Philosophy
    ## Introduction
    ...
    ## References
    Author, A. (Year). ...

Headings (## ) become APA Level-1 headings (centered, bold). A "## References"
heading starts a new page with hanging-indent entries. Requires: python-docx
(pip install python-docx).
"""
import sys, os, re
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

if len(sys.argv) < 2:
    sys.exit("usage: python scripts/make_docx.py <input.md> [output.docx]")
SRC = sys.argv[1]
OUT = sys.argv[2] if len(sys.argv) > 2 else os.path.splitext(SRC)[0] + ".docx"

with open(SRC, encoding="utf-8") as f:
    text = f.read()

# --- front-matter ---
meta = {}
fm = re.match(r"^---\s*\n(.*?)\n---\s*\n", text, re.DOTALL)
if fm:
    for line in fm.group(1).splitlines():
        if ":" in line:
            k, v = line.split(":", 1)
            meta[k.strip().lower()] = v.strip()
    text = text[fm.end():]

lines = text.splitlines()
# title: front-matter, else first "# " heading, else filename
title = meta.get("title")
for ln in lines:
    if not title and ln.startswith("# ") and not ln.startswith("## "):
        title = ln[2:].strip()
title = title or os.path.splitext(os.path.basename(SRC))[0]

doc = Document()
normal = doc.styles["Normal"]
normal.font.name = "Times New Roman"
normal.font.size = Pt(12)
normal.element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
pf = normal.paragraph_format
pf.line_spacing_rule = WD_LINE_SPACING.DOUBLE
pf.space_after = Pt(0); pf.space_before = Pt(0)
for s in doc.sections:
    s.top_margin = s.bottom_margin = Inches(1)
    s.left_margin = s.right_margin = Inches(1)

def add_page_numbers(section):
    hp = section.header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = hp.add_run()
    a = OxmlElement("w:fldChar"); a.set(qn("w:fldCharType"), "begin")
    b = OxmlElement("w:instrText"); b.set(qn("xml:space"), "preserve"); b.text = "PAGE"
    c = OxmlElement("w:fldChar"); c.set(qn("w:fldCharType"), "end")
    run._r.append(a); run._r.append(b); run._r.append(c)
    run.font.name = "Times New Roman"; run.font.size = Pt(12)
add_page_numbers(doc.sections[0])

INLINE = re.compile(r"(\*\*[^*]+\*\*|\*[^*]+\*)")
def add_inline(p, t):
    for tok in INLINE.split(t):
        if not tok:
            continue
        if tok.startswith("**") and tok.endswith("**"):
            p.add_run(tok[2:-2]).bold = True
        elif tok.startswith("*") and tok.endswith("*"):
            p.add_run(tok[1:-1]).italic = True
        else:
            p.add_run(tok)

# --- title page ---
for _ in range(3):
    doc.add_paragraph()
tp = doc.add_paragraph(); tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
tp.add_run(title).bold = True
doc.add_paragraph()
for key in ("author", "affiliation", "course", "instructor", "date"):
    if meta.get(key):
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_inline(p, meta[key])
doc.add_page_break()

bt = doc.add_paragraph(); bt.alignment = WD_ALIGN_PARAGRAPH.CENTER
bt.add_run(title).bold = True

mode = "body"; buf = []
def flush():
    global buf
    if buf:
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Inches(0.5)
        add_inline(p, " ".join(buf)); buf = []

for raw in lines:
    line = raw.rstrip()
    if line.startswith("# ") and not line.startswith("## ") and line[2:].strip() == title:
        continue
    if line.startswith("## "):
        flush()
        head = line[3:].strip()
        if head.lower() == "references":
            mode = "refs"; doc.add_page_break()
            h = doc.add_paragraph(); h.alignment = WD_ALIGN_PARAGRAPH.CENTER
            h.add_run("References").bold = True
            continue
        h = doc.add_paragraph(); h.alignment = WD_ALIGN_PARAGRAPH.CENTER
        h.add_run(head).bold = True
        continue
    if not line.strip():
        if mode == "body":
            flush()
        continue
    if mode == "refs":
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.first_line_indent = Inches(-0.5)
        add_inline(p, line.strip())
    else:
        buf.append(line.strip())
flush()

doc.save(OUT)
print("WROTE", OUT)
